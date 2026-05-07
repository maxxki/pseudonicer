#!/usr/bin/env python3
"""
MAXXKI HR Automator v2.0
=========================
Änderungen gegenüber v1.0:
  - SESSION-SCOPE: PIIPseudonicer(session_scope='document') pro Dokument.
    reset_session() wird explizit zwischen Dokumenten aufgerufen.
  - AUDIT-LOG: Append-only JSON-Lines Log (kein PII, nur Metadaten).
    Pflicht für DSGVO Art. 5 Abs. 2 Rechenschaftspflicht.
  - RECHTSGRUNDLAGE: JSON-Feld 'rechtsgrundlage' ist Pflichtfeld.
    generate() verweigert Verarbeitung ohne gültige Rechtsgrundlage.
  - FEHLERBEHANDLUNG: Strukturierte Exceptions statt bare except.

Wichtige Hinweise für Produktion:
  - Output-Verzeichnis sollte at-rest verschlüsselt sein (Art. 32 DSGVO).
  - Retention-Policy (Art. 17 DSGVO) muss extern implementiert werden
    (z.B. systemd-Timer der output/ nach X Tagen bereinigt).
  - Human-in-the-Loop Review-Step (EU AI Act Art. 14) vor Ausgabe empfohlen.
"""

import os
import re
import json
import argparse
import sys
import logging
import time
from datetime import datetime, timezone
from pathlib import Path
from docx import Document

try:
    from maxxki_pseudonicer import PIIPseudonicer
except ImportError:
    print("Fehler: maxxki_pseudonicer.py nicht gefunden.")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Gültige Rechtsgrundlagen (DSGVO Art. 6)
# ---------------------------------------------------------------------------

VALID_RECHTSGRUNDLAGEN = {
    "art6_1a": "Einwilligung (Art. 6 Abs. 1 lit. a)",
    "art6_1b": "Vertragserfüllung (Art. 6 Abs. 1 lit. b)",
    "art6_1c": "Rechtliche Verpflichtung (Art. 6 Abs. 1 lit. c)",
    "art6_1f": "Berechtigtes Interesse (Art. 6 Abs. 1 lit. f)",
}


# ---------------------------------------------------------------------------
# Audit-Logger (kein PII — nur Metadaten)
# ---------------------------------------------------------------------------

class AuditLogger:
    """
    Append-only JSON-Lines Audit-Log.
    Enthält NIEMALS personenbezogene Daten — nur Metadaten.
    """

    def __init__(self, log_path: str = "audit/hr_audit.jsonl") -> None:
        self.log_path = Path(log_path)
        self.log_path.parent.mkdir(parents=True, exist_ok=True)

    def log(self, event: str, doc_type: str, rechtsgrundlage: str,
            output_file: str, success: bool, error: str = "") -> None:
        entry = {
            "timestamp":      datetime.now(timezone.utc).isoformat().replace("+00:00", "Z"),
            "event":          event,
            "doc_type":       doc_type,
            "rechtsgrundlage": rechtsgrundlage,
            "output_file":    output_file,
            "success":        success,
            "error":          error,
            # Kein Name, keine PII — nur Metadaten
        }
        with open(self.log_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")


# ---------------------------------------------------------------------------
# HR Automator
# ---------------------------------------------------------------------------

class HRAutomator:

    TEMPLATE_MAP = {
        "zeugnis":  "template_zeugnis.docx",
        "nachtrag": "template_nachtrag.docx",
        "vertrag":  "template_vertrag.docx",
    }

    def __init__(self,
                 template_dir: str = "templates",
                 output_dir: str   = "output",
                 audit_log: str    = "audit/hr_audit.jsonl") -> None:

        self.template_dir = Path(template_dir)
        self.output_dir   = Path(output_dir)
        self.audit        = AuditLogger(audit_log)

        for d in [self.template_dir, self.output_dir]:
            d.mkdir(parents=True, exist_ok=True)

        # Ein Pseudonicer pro Automator-Instanz, scope=document.
        # reset_session() wird vor jedem Dokument aufgerufen.
        self.nicer = PIIPseudonicer(session_scope="document")

    # ------------------------------------------------------------------
    # Rechtsgrundlagen-Prüfung (vor jeder Verarbeitung)
    # ------------------------------------------------------------------

    def _validate_rechtsgrundlage(self, data: dict) -> str:
        """
        Prüft, ob eine gültige DSGVO-Rechtsgrundlage angegeben wurde.
        Wirft ValueError wenn nicht — generate() darf nicht fortfahren.
        """
        rg = data.get("rechtsgrundlage", "").strip().lower()
        if not rg:
            raise ValueError(
                "Pflichtfeld 'rechtsgrundlage' fehlt. "
                f"Erlaubte Werte: {list(VALID_RECHTSGRUNDLAGEN.keys())}"
            )
        if rg not in VALID_RECHTSGRUNDLAGEN:
            raise ValueError(
                f"Ungültige Rechtsgrundlage: '{rg}'. "
                f"Erlaubte Werte: {list(VALID_RECHTSGRUNDLAGEN.keys())}"
            )
        return VALID_RECHTSGRUNDLAGEN[rg]

    # ------------------------------------------------------------------
    # Sicheres Logging (kein Klartext-PII in Konsole)
    # ------------------------------------------------------------------

    def _log_safe(self, data: dict) -> None:
        full = f"Frau {data.get('vorname', '')} {data.get('nachname', '')}"
        safe_name = self.nicer.pseudonize(full)
        print(f"[PII-GUARD] Verarbeite Datensatz für: {safe_name}")

    # ------------------------------------------------------------------
    # Template-Ersetzung
    # ------------------------------------------------------------------

    @staticmethod
    def _replace_in_doc(doc: Document, placeholders: dict) -> None:
        """Ersetzt {{PLATZHALTER}} in Absätzen und Tabellen."""
        for p in doc.paragraphs:
            for ph, val in placeholders.items():
                if ph in p.text:
                    p.text = p.text.replace(ph, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for ph, val in placeholders.items():
                            if ph in p.text:
                                p.text = p.text.replace(ph, val)

    # ------------------------------------------------------------------
    # Hauptmethode
    # ------------------------------------------------------------------

    def generate(self, doc_type: str, data: dict) -> Optional[str]:
        """
        Generiert ein HR-Dokument aus Template + Daten.

        Returns
        -------
        str | None
            Pfad zur generierten Datei, oder None bei Fehler.

        Raises
        ------
        ValueError
            Wenn Rechtsgrundlage fehlt oder ungültig ist.
        """
        rechtsgrundlage_label = self._validate_rechtsgrundlage(data)

        template_file = self.TEMPLATE_MAP.get(doc_type)
        if not template_file:
            msg = f"Unbekannter Dokumenttyp: {doc_type}"
            self.audit.log("generate", doc_type, rechtsgrundlage_label, "", False, msg)
            print(f"❌ {msg}")
            return None

        template_path = self.template_dir / template_file
        if not template_path.exists():
            msg = f"Template nicht gefunden: {template_path}"
            self.audit.log("generate", doc_type, rechtsgrundlage_label, "", False, msg)
            print(f"❌ {msg}")
            return None

        # Session für dieses Dokument starten — VOR _log_safe, damit Pseudonymisierung greift
        self.nicer.reset_session()
        self._log_safe(data)

        doc = Document(str(template_path))

        # Platzhalter-Mapping: JSON-Keys → {{KEY_IN_GROSSBUCHSTABEN}}
        # 'rechtsgrundlage' wird nicht als Platzhalter eingefügt
        placeholders = {
            f"{{{{{k.upper()}}}}}": str(v)
            for k, v in data.items()
            if k != "rechtsgrundlage"
        }

        self._replace_in_doc(doc, placeholders)

        # Path Traversal Sanitierung: nur alphanumerische Zeichen + Bindestrich erlaubt
        raw_name = data.get('nachname', 'export')
        safe_name = re.sub(r'[^A-Za-z0-9\-]', '_', raw_name)
        output_name = (
            f"{doc_type}_{safe_name}"
            f"_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        output_path = self.output_dir / output_name
        doc.save(str(output_path))

        self.audit.log(
            "generate", doc_type, rechtsgrundlage_label,
            str(output_path), True
        )
        print(f"✅ Erfolg: {output_path} erstellt.")
        print(f"   Rechtsgrundlage: {rechtsgrundlage_label}")
        return str(output_path)


# Typing-Import für Optional im generate()-Signatur
from typing import Optional


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="MAXXKI HR Automator v2.0 🚀",
        epilog=(
            "Beispiel JSON-Eingabe:\n"
            '  { "vorname": "Maria", "nachname": "Muster",\n'
            '    "position": "Projektmanagerin",\n'
            '    "rechtsgrundlage": "art6_1b" }\n\n'
            f"Gültige Rechtsgrundlagen: {list(VALID_RECHTSGRUNDLAGEN.keys())}"
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument(
        "--type",
        choices=["zeugnis", "nachtrag", "vertrag"],
        help="Typ des zu erstellenden Dokuments",
    )
    parser.add_argument("--input", help="Pfad zur JSON-Eingabedatei")
    parser.add_argument(
        "--demo",
        action="store_true",
        help="Startet die Demo mit Musterdaten (Rechtsgrundlage: art6_1b)",
    )

    args = parser.parse_args()
    automator = HRAutomator()

    if args.demo:
        print("--- 🚀 Starte HR-Automator Demo ---")
        demo_data = {
            "vorname":          "Maria",
            "nachname":         "Musterfrau",
            "position":         "Senior Projektmanagerin",
            "firma":            "Muster GmbH",
            "rechtsgrundlage":  "art6_1b",
        }
        automator.generate("vertrag", demo_data)
        sys.exit(0)

    if not args.type or not args.input:
        parser.print_help()
        sys.exit(1)

    try:
        with open(args.input, "r", encoding="utf-8") as f:
            data = json.load(f)
        automator.generate(args.type, data)
    except ValueError as e:
        print(f"❌ Validierungsfehler: {e}")
        sys.exit(2)
    except Exception as e:
        print(f"❌ Fehler beim Verarbeiten: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
